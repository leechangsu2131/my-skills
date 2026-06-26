import docx
from pathlib import Path

def main():
    docx_path = Path("skills/classmanage-evaluate-to-neis/data/2026학년도 3학년 1학기 행동특성 및 종합의견(수정본).docx")
    if not docx_path.exists():
        print(f"Error: {docx_path} does not exist.")
        return
        
    doc = docx.Document(docx_path)
    output_lines = []
    
    output_lines.append(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    # Process paragraphs
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            output_lines.append(f"[{idx}] {text}")
            
    output_lines.append("\nTables:")
    for idx, table in enumerate(doc.tables):
        output_lines.append(f"Table {idx}: {len(table.rows)} rows, {len(table.columns)} columns")
        for r_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            output_lines.append(f"  Row {r_idx}: {cells}")

    out_file = Path("scratch/parsed_docx_content.txt")
    out_file.write_text("\n".join(output_lines), encoding="utf-8")
    print(f"Wrote parsed content to {out_file}")

if __name__ == "__main__":
    main()
