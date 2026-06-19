#!/usr/bin/env python
"""Generate class operating fund evidence pages as a DOCX file."""

from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


def money(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        return f"{int(str(value).replace(',', '').strip()):,}"
    except ValueError:
        return str(value)


def load_data(path: Path) -> dict[str, Any]:
    if path.suffix.lower() == ".json":
        return json.loads(path.read_text(encoding="utf-8-sig"))

    if path.suffix.lower() == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            rows = list(csv.DictReader(f))
        receipts: list[dict[str, Any]] = []
        for row in rows:
            receipts.append(
                {
                    "date": row.get("date") or row.get("사용일자") or "",
                    "vendor": row.get("vendor") or row.get("사용업체명") or "",
                    "activity": row.get("activity") or row.get("활동구분") or "",
                    "usage": row.get("usage") or row.get("사용내역") or "",
                    "amount": row.get("amount") or row.get("사용금액") or "",
                    "evidence_type": row.get("evidence_type") or row.get("증빙구분") or "",
                    "items": [
                        {
                            "name": row.get("item") or row.get("품목") or row.get("usage") or "",
                            "quantity": row.get("quantity") or row.get("수량") or "",
                            "amount": row.get("item_amount") or row.get("금액") or row.get("amount") or "",
                        }
                    ],
                }
            )
        return {"receipts": receipts}

    raise ValueError("Input must be .json or .csv")


def set_font(paragraph, size: int = 10, bold: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "맑은 고딕"
        run.font.size = Pt(size)
        run.bold = bold


def add_label(paragraph, label: str, value: Any) -> None:
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(str(value or ""))
    set_font(paragraph, 10)


def add_receipt_page(doc: Document, data: dict[str, Any], receipt: dict[str, Any], index: int) -> None:
    if index > 0:
        doc.add_section(WD_SECTION_START.NEW_PAGE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = (
        f"{data.get('school_year', 2026)}학년도 "
        f"{data.get('grade', '')}학년 {data.get('class_no', '')}반 "
        "학급 자율 운영비 집행 증빙자료"
    )
    title.add_run(title_text).bold = True
    set_font(title, 15, True)

    meta = doc.add_table(rows=3, cols=4)
    meta.style = "Table Grid"
    values = [
        ("사용일자", receipt.get("date", ""), "사용업체명", receipt.get("vendor", "")),
        ("사용금액", money(receipt.get("amount", "")), "증빙구분", receipt.get("evidence_type", "")),
        ("활동 구분", receipt.get("activity", ""), "담임", data.get("teacher", "")),
    ]
    for row, row_values in zip(meta.rows, values):
        for cell, text in zip(row.cells, row_values):
            cell.text = str(text)
            for p in cell.paragraphs:
                set_font(p, 9)

    doc.add_paragraph()
    usage = doc.add_paragraph()
    add_label(usage, "사용내역", receipt.get("usage", ""))

    doc.add_paragraph("영수증 또는 전산자료 부착")
    receipt_box = doc.add_table(rows=1, cols=1)
    receipt_box.style = "Table Grid"
    box_cell = receipt_box.cell(0, 0)
    box_cell.text = "\n\n\n\n\n\n\n"
    note = box_cell.paragraphs[0]
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("한 면에 내용이 모두 보이도록 겹치지 않게 첨부")
    set_font(note, 10)

    doc.add_paragraph("구매 내역")
    items = receipt.get("items") or []
    item_table = doc.add_table(rows=1, cols=3)
    item_table.style = "Table Grid"
    headers = ["품목", "수량", "총액(원)"]
    for cell, header in zip(item_table.rows[0].cells, headers):
        cell.text = header
        for p in cell.paragraphs:
            set_font(p, 9, True)

    if not items:
        items = [{"name": "", "quantity": "", "amount": receipt.get("amount", "")}]

    for item in items:
        row = item_table.add_row().cells
        row[0].text = str(item.get("name", ""))
        row[1].text = str(item.get("quantity", ""))
        row[2].text = money(item.get("amount", ""))
        for cell in row:
            for p in cell.paragraphs:
                set_font(p, 9)

    footer = doc.add_paragraph()
    footer.add_run("제출 유의사항: ").bold = True
    footer.add_run("정산서를 표지로 사용하고, 증빙자료는 날짜순으로 부착합니다. 카드영수증은 복사본을 함께 첨부합니다.")
    set_font(footer, 9)


def build_doc(data: dict[str, Any]) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    receipts = data.get("receipts") or []
    if not receipts:
        raise ValueError("No receipts found. Provide a receipts array or CSV rows.")

    receipts = sorted(receipts, key=lambda r: str(r.get("date", "")))
    for index, receipt in enumerate(receipts):
        add_receipt_page(doc, data, receipt, index)
    return doc


def main() -> None:
    parser = argparse.ArgumentParser(description="Generate 학급 자율 운영비 증빙자료 DOCX pages.")
    parser.add_argument("input", type=Path, help="JSON or CSV input")
    parser.add_argument("--output", "-o", type=Path, default=Path("증빙자료.docx"))
    args = parser.parse_args()

    data = load_data(args.input)
    doc = build_doc(data)
    doc.save(args.output)
    print(f"saved: {args.output}")


if __name__ == "__main__":
    main()
