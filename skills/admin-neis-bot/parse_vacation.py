#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""여름방학교육계획 docx 파일 파싱"""
import docx
from pathlib import Path

def main():
    docx_path = Path(__file__).parent / "data" / "26학년도 여름방학교육계획.docx"
    if not docx_path.exists():
        print(f"Error: {docx_path} does not exist.")
        return
    
    doc = docx.Document(docx_path)
    lines = []
    
    lines.append(f"Number of paragraphs: {len(doc.paragraphs)}")
    for idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            lines.append(f"[{idx}] {text}")
    
    lines.append("\n---TABLES---")
    for ti, table in enumerate(doc.tables):
        lines.append(f"Table {ti}: {len(table.rows)} rows, {len(table.columns)} columns")
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip().replace('\n', ' | ') for cell in row.cells]
            lines.append(f"  Row {ri}: {cells}")
    
    out_file = Path(__file__).parent / "data" / "parsed_vacation_plan.txt"
    out_file.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {len(lines)} lines to {out_file}")

if __name__ == "__main__":
    main()
